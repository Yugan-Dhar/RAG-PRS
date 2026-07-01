from typing import List, Dict, Any
import os
from .base import BaseParser

class DocxParser(BaseParser):
    def __init__(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            self.docx = None

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        if not self.docx:
            raise RuntimeError("python-docx is not installed")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        doc = self.docx.Document(file_path)
        elements = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                elements.append({
                    "text": text,
                    "metadata": {
                        "type": "paragraph",
                        "style": para.style.name if para.style else "Normal",
                        "index": i
                    }
                })
                
        # Basic table parsing
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(filter(None, row_data))
                if row_text:
                    elements.append({
                        "text": row_text,
                        "metadata": {
                            "type": "table_row"
                        }
                    })
                    
        return elements
